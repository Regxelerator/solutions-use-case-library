import json
from pathlib import Path
from typing import Union, Dict, Any, List
import copy

from services.llm_client import LLMClientFactory
from utils.helpers import log_to_csv
from concurrent.futures import ThreadPoolExecutor, as_completed
from tqdm import tqdm

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from utils.schemas import MODEL_INPUT_SCHEMA

INPUT_DIR = Path("input")
REQ_FILE  = INPUT_DIR / "document_requirements.json"
PROMPT_DIR = Path("prompts")
PROMPT_VALIDATION_FILE = PROMPT_DIR / "PROMPT_VALIDATION.txt"


class ValidationProcessor:
    """
    • pulls *<Entity> - Extracted Documents* folders
    • builds / uploads document_list.json
    • validates against regulatory checklist
    """

    def __init__(self, ms_graph_client, output_drive_id):
        self.output_drive_id = output_drive_id
        self.ms_graph_client    = ms_graph_client
        self.llm_factory        = LLMClientFactory()

    def write_json_file(self, path: Union[str, Path], data: Union[dict, list], indent: int = 2) -> None:
        with open(path, "w", encoding="utf-8") as fp:
            json.dump(data, fp, ensure_ascii=False, indent=indent)

    def load_json_file(self, path: Union[str, Path]) -> Union[dict, list]:
        with open(path, "r", encoding="utf-8") as fp:
            return json.load(fp)

    def upload_json_to_sharepoint(self, folder_name: str, json_name: str, data: Union[dict, list]) -> None:
        target_path = f"{folder_name}/{json_name}"
        url = f"https://graph.microsoft.com/v1.0/drives/{self.output_drive_id}/root:/{target_path}:/content"
        payload = json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")
        resp = self.ms_graph_client.put(url, payload) 
        resp.raise_for_status()

    @staticmethod
    def get_file_content_and_meta_data(_ms_graph_client, output_drive_id, folder, subfolder):
        subfolder_path = f"{folder}/{subfolder}"
        folder_url = (
            f"https://graph.microsoft.com/v1.0/drives/{output_drive_id}"
            f"/root:/{subfolder_path}:/children"
        )
        res_folder = _ms_graph_client.get(folder_url)
        res_folder.raise_for_status()

        folder_contents = res_folder.json().get("value", [])
        if not folder_contents:
            print(f"\nFolder '{subfolder_path}' is empty. Skipping…")
            return {}, {}

        content_data, meta_data = {}, {}
        for child in folder_contents:
            child_name = child.get("name")
            if child_name in {"content.json", "metadata.json"}:
                file_id = child.get("id")
                url = (
                    f"https://graph.microsoft.com/v1.0/drives/"
                    f"{output_drive_id}/items/{file_id}/content"
                )
                resp = _ms_graph_client.get(url)
                resp.raise_for_status()
                data = json.loads(resp.content.decode("utf-8"))
                if child_name == "content.json":
                    content_data = data
                else:
                    meta_data = data
        return content_data, meta_data

    def extract_target_folder_and_files(self, entity_name: str | None = None) -> Dict[str, Dict[str, Any]]:
        try:
            url_root = f"https://graph.microsoft.com/v1.0/drives/{self.output_drive_id}/root/children"
            res_root = self.ms_graph_client.get(url_root)
            res_root.raise_for_status()
            items = res_root.json().get("value", [])

            combined: Dict[str, Dict[str, Any]] = {}

            def should_process(name: str) -> bool:
                if entity_name:
                    return name.lower().startswith(entity_name.lower()) and name.lower().endswith(" - extracted documents")
                return name.lower().endswith(" - extracted documents")

            def process_folder(item):
                folder_name = item.get("name")
                folder_url = f"https://graph.microsoft.com/v1.0/drives/{self.output_drive_id}/root:/{folder_name}:/children"
                try:
                    res_folder = self.ms_graph_client.get(folder_url)
                    res_folder.raise_for_status()
                    folder_contents = res_folder.json().get("value", [])
                    if not folder_contents:
                        print(f"\n Folder '{folder_name}' is empty. Skipping…")
                        return

                    file_map: Dict[str, Dict[str, Any]] = {}
                    for child in folder_contents:
                        if "folder" not in child:          
                            continue
                        
                        content, meta = self.get_file_content_and_meta_data(
                            self.ms_graph_client,
                            self.output_drive_id,
                            folder_name,
                            child.get("name"),
                        )
                        meta["content"] = content
                        meta["name"]    = child.get("name")
                        file_map[child.get("name")] = meta

                    combined[folder_name] = file_map
                except Exception as exc:
                    print(f"Exception reading '{folder_name}': {exc}")

            candidates = [i for i in items if i.get("folder") and should_process(i["name"])]

            with ThreadPoolExecutor(max_workers=4) as executor:
                futures = {executor.submit(process_folder, itm): itm["name"] for itm in candidates}
                for fut in tqdm(
                    as_completed(futures),
                    total=len(futures),
                    ncols=100,
                    bar_format="{l_bar}{bar} | {n_fmt}/{total_fmt} • ETA: {remaining}",
                    desc="\n Processing",
                    unit="folder",
                ):
                    try:
                        fut.result()
                    except Exception as err:
                        print(f"\n Error processing folder '{futures[fut]}': {err}")

            return combined
        except Exception as e:
            raise Exception(f"Exception reading drive folders: {e}")

    def convert_file_map_to_doc_list(self, dir_map: dict) -> list:
        """
        Convert the nested {folder: {file: metadata}} structure returned from
        extract_target_folder_and_files into the flat list expected by the LLM.
        """
        documents: list = []
        for _, files in dir_map.items():
            for _, meta in files.items():
                content   = meta.get("content", {})
                documents.append(
                    {
                        "document_name":    meta.get("name"),
                        "document_type":    meta.get("document_type")    or content.get("document_type"),
                        "document_summary": meta.get("document_summary") or content.get("document_summary"),
                        "document_toc":     meta.get("document_toc")     or content.get("document_toc"),
                    }
                )
        return documents

    @staticmethod
    def _load_requirements() -> List[str]:
        with REQ_FILE.open(encoding="utf-8") as fp:
            data = json.load(fp)
        return data if isinstance(data, list) else data["DOCUMENT_REQUIREMENTS"]
    
    def load_document_requirements(self) -> List[str]:

        return self._load_requirements()


    @staticmethod
    def _build_validation_schema(checklist: List[str]) -> Dict[str, Any]:
        schema = copy.deepcopy(MODEL_INPUT_SCHEMA["validation"])
        enum_node = (
            schema["schema"]["properties"]["reviews"]["items"]
                  ["properties"]["document_type"]
        )
        enum_node["enum"] = checklist
        return schema

    def validate_document_list(self,
                               checklist: List[str],
                               document_list: List[dict],
    ) -> dict:
        schema_obj = self._build_validation_schema(checklist)

        user_message = {
            "document_checklist": checklist,
            "document_list":      document_list,
        }

        with PROMPT_VALIDATION_FILE.open(encoding="utf-8") as f:
            system_prompt = f.read().strip()

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": json.dumps(user_message, ensure_ascii=False)},
        ]

        try:
            log_to_csv(f"== DOCUMENT CHECKLIST LLM INPUT MESSAGE ==\n{messages}")
            response = self.llm_factory.chat_completion(
                "validation",
                messages,
                json_schema=schema_obj,
            )

            if isinstance(response, str):
                response = json.loads(response)
            if isinstance(response, dict):
                return response
            raise ValueError("LLM response is not a JSON object.")
        except Exception as exc:
            log_to_csv(f"== DOCUMENT CHECKLIST LLM EXCEPTION ==\n{exc}")
            raise
    
    @staticmethod
    def _set_cell_shading(cell, fill_color: str) -> None:

        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_color)
        cell._tc.get_or_add_tcPr().append(shd)

    @staticmethod
    def _set_cell_border(cell, size: int = 4, color: str = "000000") -> None:

        tc_pr = cell._tc.get_or_add_tcPr()
        for pos in ("top", "left", "bottom", "right"):
            ln = OxmlElement(f"w:{pos}")
            ln.set(qn("w:val"), "single")
            ln.set(qn("w:sz"), str(size))
            ln.set(qn("w:color"), color)
            ln.set(qn("w:space"), "0")
            tc_pr.append(ln)

    def json_to_word(
        self,
        json_path: Union[str, Path],
        word_path: Union[str, Path],
    ) -> None:

        rows = self.load_json_file(json_path)
        if isinstance(rows, dict) and "reviews" in rows:
            rows = rows["reviews"]

        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = (
            section.page_height,
            section.page_width,
        )
        for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
            setattr(section, m, Inches(0.7))

        heading = doc.add_heading("Document Validation Result", level=1)
        heading.alignment = 0 
        run = heading.runs[0]
        run.font.name = "Calibri"           
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 0, 0)  

        doc.add_paragraph() 

        table = doc.add_table(rows=len(rows) + 1, cols=4)
        table.autofit = False

        table.columns[0].width = Inches(0.5)
        table.columns[1].width = Inches(4.0)
        table.columns[2].width = Inches(1.0)
        table.columns[3].width = Inches(4.1)

        headers = ["#", "Document Required", "Validation Status", "Comments"]
        for col_idx, txt in enumerate(headers):
            cell = table.cell(0, col_idx)
            para = cell.paragraphs[0]
            run = para.add_run(txt)
            run.font.name = "Calibri"
            run.font.bold = True
            run.font.size = Pt(10)
            self._set_cell_shading(cell, "E8E8E8")
            self._set_cell_border(cell)

        rating_colour = {
            "Not satisfied":      "C00000",
            "Fully satisfied":    "4EA72E",
            "Partially satisfied": "FFC000",
        }

        for row_idx, record in enumerate(rows, start=1):
            values = [
                str(row_idx),
                record.get("document_type", ""),
                record.get("validation_result", ""),
                record.get("comments", ""),
            ]
            for col_idx, val in enumerate(values):
                cell = table.cell(row_idx, col_idx)
                para = cell.paragraphs[0]
                run = para.add_run(val)
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                self._set_cell_border(cell)

                if col_idx == 2:
                    self._set_cell_shading(cell, rating_colour.get(val, "FFFFFF"))

        doc.save(word_path)
