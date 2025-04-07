import json
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from agents import function_tool
import os
from typing import Dict


def write_to_json_file(file_path: str, data, indent: int = 4) -> None:
    """
    Writes data to a JSON file.

    Args:
        file_path (str): Path to the output JSON file.
        data (dict): Data to be written into the file.
        indent (int, optional): Number of spaces for indentation in JSON. Defaults to 4.

    Returns:
        None
    """
    with open(file_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=indent)


def read_json_file(file_path: str) -> dict:
    """
    Reads and loads data from a JSON file.

    Args:
        file_path (str): Path to the JSON file.

    Returns:
        dict: Parsed JSON data.
    """
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError) as e:
        print(f"Error reading JSON file '{file_path}': {e}")
        return {}


@function_tool
async def create_report(output_dir: str) -> Dict[str, str]:
    """
    Generates a Word document report.

    Args:
        output_dir (str): The directory where the final report and related files will be saved.

    Returns:
        Dict[str, str]: A dictionary indicating the status of the report generation and the generated filename.
    """

    def set_run_font(run, font_name="Calibri", font_size=10, is_bold=False):
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = is_bold
        run.font.color.rgb = RGBColor(0, 0, 0)

    def set_heading_run_font(
        run, font_name="Calibri Light", font_size=14, is_bold=True
    ):
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = is_bold
        run.font.color.rgb = RGBColor(0, 0, 0)

    def create_footer_disclaimer_and_page_number(section):
        footer = section.footer
        para_left = footer.add_paragraph()
        para_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_left = para_left.add_run(
            "Disclaimer: This output was created through the application of generative AI. "
            "Please validate accuracy and completeness before using it for decision-making."
        )
        set_run_font(run_left, font_size=8, font_name="Calibri")

    def generate_word_report(data, final_world_report="Final_Report.docx"):
        doc = Document()

        title_section = doc.sections[0]
        title_section.left_margin = Inches(1)
        title_section.right_margin = Inches(1)
        title_section.top_margin = Inches(4)
        title_section.bottom_margin = Inches(3)

        title_par = doc.add_paragraph()
        title_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        Report_Title = "AI Supervision in Financial Services"
        run_title = title_par.add_run(Report_Title)
        set_heading_run_font(
            run_title, font_size=28, font_name="Calibri Light", is_bold=True
        )

        current_month_year = datetime.now().strftime("%B %Y")
        date_par = doc.add_paragraph()
        date_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_date = date_par.add_run(current_month_year)
        set_run_font(run_date, font_size=14, font_name="Calibri")

        doc.add_page_break()

        content_section = doc.add_section()
        content_section.left_margin = Inches(1)
        content_section.right_margin = Inches(1)
        content_section.top_margin = Inches(1)
        content_section.bottom_margin = Inches(1.5)

        create_footer_disclaimer_and_page_number(content_section)

        is_first_chapter = True
        for item in data["report"]:
            chapter = item.get("response", {})
            chapter_title = chapter.get("title", "Untitled Chapter")
            subsections = chapter.get("subsections", [])

            chapter_heading = doc.add_heading(chapter_title, level=1)
            chapter_heading.style = doc.styles["Heading 1"]
            set_run_font(
                chapter_heading.runs[0],
                font_name="Calibri Light",
                font_size=14,
                is_bold=True,
            )

            if not is_first_chapter:
                chapter_heading.paragraph_format.page_break_before = True
            is_first_chapter = False

            for subsection in subsections:
                sub_title = subsection.get("title", "Untitled Subsection")
                sub_text = subsection.get("text", "")

                sub_heading = doc.add_heading(sub_title, level=2)
                sub_heading.style = doc.styles["Heading 2"]
                set_run_font(
                    sub_heading.runs[0], font_name="Calibri", font_size=12, is_bold=True
                )

                text_par = doc.add_paragraph()
                text_run = text_par.add_run(sub_text.replace("\n", " "))
                set_run_font(text_run, font_size=10, font_name="Calibri")

                text_par.paragraph_format.space_after = Pt(8)

        doc.save(final_world_report)
        print(f"Report saved as {final_world_report}")

    final_report_json = os.path.join(output_dir, "final_report.json")

    report_data = read_json_file(final_report_json)
    final_world_report = os.path.join(output_dir, "Final Report", "Final_Report.docx")
    os.makedirs(os.path.dirname(final_world_report), exist_ok=True)
    generate_word_report(report_data, final_world_report)
    return {"status": "Report generated successfully", "filename": "Final_Report.docx"}
