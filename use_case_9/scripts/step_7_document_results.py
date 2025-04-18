# scripts/generate_risk_comparative_report.py

import os
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_run_font(run, font_name="Calibri", font_size=10, is_bold=False, font_color="000000"):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = is_bold
    run.font.color.rgb = RGBColor.from_string(font_color)


def set_heading_run_font(run, font_name="Calibri Light", font_size=14, is_bold=True, font_color="000000"):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = is_bold
    run.font.color.rgb = RGBColor.from_string(font_color)


def create_footer_disclaimer_and_page_number(section):
    footer = section.footer
    footer.paragraphs.clear()

    p_left = footer.add_paragraph()
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_left = p_left.add_run(
        "Disclaimer: This output was created through the application of generative AI. "
        "Please validate accuracy and completeness before using it for decision‑making."
    )
    set_run_font(r_left, font_size=8)

    p_right = footer.add_paragraph()
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_right = p_right.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    for elm in (fld_begin, instr, fld_sep, fld_end):
        r_right._r.append(elm)
    set_run_font(r_right, font_size=8)


def set_cell_shading(cell, fill_color):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), fill_color)
    cell._tc.get_or_add_tcPr().append(sh)


def set_cell_borders(cell, border_color="000000", border_size="4"):
    tcPr = cell._element.get_or_add_tcPr()
    for side in ("top", "left", "bottom", "right"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn("w:val"), "single")
        bd.set(qn("w:sz"), border_size)
        bd.set(qn("w:color"), border_color)
        bd.set(qn("w:space"), "0")
        tcPr.append(bd)


def generate_risk_comparative_report(output_root_dir: str, final_report_path: str):
    """
    Reads:
      - output/<year>/interim_outputs/synthesis_annual_report_<idx>.json  (for two years)
      - output/comparative_analysis/comparative_analysis_annual_report_<idx>.json
    For each <idx>, builds a landscape Word doc section with:
      * Sub‑header showing the actual bank_name from the JSON
      * A 4‑column table: [Risk Dimension, Year1, Year2, Comparative]
    Saves the document to final_report_path.
    """
    # find the two year folders (excluding the comparative_analysis folder)
    years = sorted(
        d for d in os.listdir(output_root_dir)
        if os.path.isdir(os.path.join(output_root_dir, d)) and d != "comparative_analysis"
    )
    comp_dir = os.path.join(output_root_dir, "comparative_analysis")
    if len(years) < 2 or not os.path.isdir(comp_dir):
        print("Need two years of synthesis + comparative_analysis folder.")
        return

    comp_files = sorted(
        f for f in os.listdir(comp_dir)
        if f.startswith("comparative_analysis_annual_report_") and f.endswith(".json")
    )
    if not comp_files:
        print("No comparative files found.")
        return

    # prepare document
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    create_footer_disclaimer_and_page_number(section)

    # overall title
    title_par = doc.add_paragraph()
    title_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = title_par.add_run("Comparative Risk Analysis Report")
    set_heading_run_font(run_title, font_size=20)
    doc.add_paragraph()

    # keys now use underscores, nested under "synthesis"
    dimensions = ["risk_exposure", "risk_controls", "risk_incidents", "risk_governance"]

    def unwrap_name(obj):
        """Turn ['Bank Name'] into 'Bank Name' (or return empty string)."""
        name = obj.get("bank_name")
        if isinstance(name, list) and name:
            return name[0]
        if isinstance(name, str):
            return name
        return ""

    for comp_fname in comp_files:
        idx = comp_fname.replace("comparative_analysis_annual_report_", "").replace(".json", "")

        # load the two years' synthesis JSON
        year1_path = os.path.join(output_root_dir, years[0], "interim_outputs",
                                  f"synthesis_annual_report_{idx}.json")
        year2_path = os.path.join(output_root_dir, years[1], "interim_outputs",
                                  f"synthesis_annual_report_{idx}.json")
        comp_path  = os.path.join(comp_dir, comp_fname)

        with open(year1_path, encoding="utf-8") as f1:
            data1 = json.load(f1)
        with open(year2_path, encoding="utf-8") as f2:
            data2 = json.load(f2)
        with open(comp_path, encoding="utf-8") as fc:
            comp = json.load(fc)

        # determine bank_name (unwrap from list)
        bank_name = unwrap_name(data1) or unwrap_name(data2) or unwrap_name(comp) or f"Bank {idx}"

        # get the nested syntheses
        y1 = data1.get("synthesis", {})
        y2 = data2.get("synthesis", {})
        yc = comp.get("synthesis", {})

        # sub‑header
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rh = h.add_run(bank_name)
        set_heading_run_font(rh, font_size=16)
        doc.add_paragraph()

        # table setup
        table = doc.add_table(rows=1, cols=4)
        table.autofit = True
        hdr = table.rows[0].cells
        col_headers = ["Risk Dimension", years[0], years[1], "Comparative Analysis"]
        for i, text in enumerate(col_headers):
            hdr[i].text = ""
            r = hdr[i].paragraphs[0].add_run(text)
            set_run_font(r, font_size=10, is_bold=True, font_color="FFFFFF")
            set_cell_shading(hdr[i], "000000")
            set_cell_borders(hdr[i])

        # fill rows
        for dim in dimensions:
            row = table.add_row().cells

            # label: e.g. "Risk Exposure"
            label = dim.replace("_", " ").title()
            run_d = row[0].paragraphs[0].add_run(label)
            set_run_font(run_d, is_bold=True)
            set_cell_shading(row[0], "D9D9D9")
            set_cell_borders(row[0])

            # helper to turn list→newline or str
            def render(cell_data):
                if isinstance(cell_data, list):
                    return "\n".join(cell_data)
                return str(cell_data or "")

            # Year 1
            run1 = row[1].paragraphs[0].add_run(render(y1.get(dim)))
            set_run_font(run1)
            set_cell_borders(row[1])

            # Year 2
            run2 = row[2].paragraphs[0].add_run(render(y2.get(dim)))
            set_run_font(run2)
            set_cell_borders(row[2])

            # Comparative
            runc = row[3].paragraphs[0].add_run(render(yc.get(dim)))
            set_run_font(runc)
            set_cell_borders(row[3])

        doc.add_page_break()

    # save final document
    os.makedirs(os.path.dirname(final_report_path), exist_ok=True)
    doc.save(final_report_path)
    print(f"Report saved → {final_report_path}")
    return final_report_path