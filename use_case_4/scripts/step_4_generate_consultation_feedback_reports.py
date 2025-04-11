import sys
import os
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

import pandas as pd
from datetime import datetime
from collections import defaultdict
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from utils.file_handler import load_json


def set_run_font(run, font_name="Calibri", font_size=10, is_bold=False):
    """
    Sets the font style for a given text run in a Word document.

    Args:
        run: A `Run` object from `python-docx` representing a segment of text.
        font_name (str, optional): The name of the font to apply. Defaults to "Calibri".
        font_size (int, optional): The font size in points. Defaults to 10.
        is_bold (bool, optional): Whether the text should be bold. Defaults to False.
    """
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = is_bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_heading_run_font(run, font_name="Calibri Light", font_size=14, is_bold=True):
    """
    Sets the font style for a heading run in a Word document.

    Args:
        run: A `Run` object from `python-docx` representing a heading text segment.
        font_name (str, optional): The name of the font to apply. Defaults to "Calibri Light".
        font_size (int, optional): The font size in points. Defaults to 14.
        is_bold (bool, optional): Whether the heading should be bold. Defaults to True.
    """
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = is_bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def insert_table_of_contents(doc):
    """
    Inserts a Table of Contents (ToC) into a Word document.

    Args:
        doc: A `Document` object from `python-docx` representing the Word document.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar_begin)
    run._r.append(fldChar_separate)
    run._r.append(fldChar_end)


def create_footer_disclaimer_and_page_number(section):
    """
    Adds a footer to a Word document section with a disclaimer and page number.

    Args:
        section: A `Section` object from `python-docx` representing the document section.
    """
    footer = section.footer
    footer.paragraphs.clear()

    para_left = footer.add_paragraph()
    para_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_left = para_left.add_run(
        "Disclaimer: This output was created through the application of generative AI. "
        "Please validate accuracy and completeness before using it for decision-making."
    )
    set_run_font(run_left, font_size=8, font_name="Calibri")

    para_right = footer.add_paragraph()
    para_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run_right = para_right.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"

    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    run_right._r.append(fldChar_begin)
    run_right._r.append(instrText)
    run_right._r.append(fldChar_separate)
    run_right._r.append(fldChar_end)

    set_run_font(run_right, font_size=8, font_name="Calibri")


def sort_key(ch):
    """
    Converts a string to an integer if possible, otherwise assigns a high default value.

    Args:
        ch: A string representing a potential number.

    Returns:
        int: The integer representation of the string if conversion succeeds; otherwise, returns a large default value (999999).

    """
    try:
        return int(ch)
    except ValueError:
        return 999999


def main_word(
    conult_quest_summary_json: str,
    consult_paper_info_json: str,
    executive_summary_json: str,
    output_dir: str,
) -> None:
    """
    Generates a Word document summarizing consultation questions and responses.

    Args:
        conult_quest_summary_json (str): Path to consultation question summaries JSON file.
        consult_paper_info_json (str): Path to consultation paper information JSON file.
        executive_summary_json (str): Path to executive summary JSON file.
        output_dir (str): Directory path where the Word document will be saved.
    """
    summaries = load_json(conult_quest_summary_json)
    consultation_paper = load_json(consult_paper_info_json)
    executive_summary_data = load_json(executive_summary_json)

    consultation_title = consultation_paper.get("General_Information", {}).get(
        "Consultation_Title", "Untitled Consultation"
    )

    chapters_map = defaultdict(list)
    for item in summaries:
        chap_id = item["chapter_id"]
        chapters_map[chap_id].append(item)

    sorted_chapters = sorted(chapters_map.keys(), key=sort_key)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(3)
    section.bottom_margin = Inches(3)

    create_footer_disclaimer_and_page_number(section)

    title_par_1 = doc.add_paragraph()
    title_par_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title_1 = title_par_1.add_run("Summary Report")
    set_heading_run_font(
        run_title_1, font_size=28, font_name="Calibri Light", is_bold=True
    )

    title_par_2 = doc.add_paragraph()
    title_par_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title_2 = title_par_2.add_run(f"Consultation: {consultation_title}")
    set_heading_run_font(
        run_title_2, font_size=20, font_name="Calibri Light", is_bold=False
    )

    current_month_year = datetime.now().strftime("%B %Y")
    date_par = doc.add_paragraph()
    date_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = date_par.add_run(current_month_year)
    set_run_font(run_date, font_size=14, font_name="Calibri")

    doc.add_page_break()

    exec_summary_heading = doc.add_heading("Executive Summary", level=1)
    exec_summary_heading.style = doc.styles["Heading 1"]
    set_run_font(
        exec_summary_heading.runs[0],
        font_name="Calibri Light",
        font_size=14,
        is_bold=True,
    )

    exec_summary_par = doc.add_paragraph()
    exec_summary_run = exec_summary_par.add_run(
        executive_summary_data.get("executive_summary", "")
    )
    set_run_font(exec_summary_run, font_size=10, font_name="Calibri")

    next_section = doc.sections[-1]
    next_section.left_margin = Inches(1)
    next_section.right_margin = Inches(1)
    next_section.top_margin = Inches(1)
    next_section.bottom_margin = Inches(1)

    is_first_chapter = True
    for chap_id in sorted_chapters:
        first_item = chapters_map[chap_id][0]
        chapter_title = first_item["chapter_title"].replace("#", "").strip()

        chapter_heading = doc.add_heading(chapter_title, level=1)
        chapter_heading.style = doc.styles["Heading 1"]
        chapter_heading.add_run(chapter_title)
        set_run_font(
            chapter_heading.runs[0],
            font_name="Calibri Light",
            font_size=14,
            is_bold=True,
        )

        if not is_first_chapter:
            chapter_heading.paragraph_format.page_break_before = True
        is_first_chapter = False

        for item in chapters_map[chap_id]:
            q_heading = doc.add_heading("", level=2)
            q_heading.style = doc.styles["Heading 2"]
            q_heading.paragraph_format.space_after = Pt(8)
            q_heading.paragraph_format.left_indent = Inches(0.1)
            q_heading.paragraph_format.right_indent = Inches(0.1)

            run_q = q_heading.add_run(item["question"])
            set_run_font(run_q, font_size=11, font_name="Calibri", is_bold=True)

            pPr = q_heading._p.get_or_add_pPr()

            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F2F2F2")
            pPr.append(shd)

            pBdr = OxmlElement("w:pBdr")
            for side in ["top", "left", "bottom", "right"]:
                border_elm = OxmlElement(f"w:{side}")
                border_elm.set(qn("w:val"), "single")
                border_elm.set(qn("w:sz"), "6")
                border_elm.set(qn("w:space"), "6")
                border_elm.set(qn("w:color"), "808080")
                pBdr.append(border_elm)
            pPr.append(pBdr)

            summary_par = doc.add_paragraph()
            run_summ = summary_par.add_run(item["response_summary"])
            set_run_font(run_summ, font_size=10, font_name="Calibri")
            summary_par.paragraph_format.space_after = Pt(8)

    output_filename = os.path.join(output_dir, "Consultation_Summary_Report.docx")
    doc.save(output_filename)


def consolidating_the_results(
    consult_quest_summary_json: str,
    consult_paper_info_json: str,
    executive_summary_json: str,
    consultation_quest_json: str,
    segmented_json_file: str,
    output_dir: str,
) -> None:
    """
    Consolidates consultation results by generating Word and Excel reports.

    Args:
        consult_quest_summary_json (str): Path to consultation question summaries JSON file.
        consult_paper_info_json (str): Path to consultation paper information JSON file.
        executive_summary_json (str): Path to executive summary JSON file.
        consultation_quest_json (str): Path to consultation questions JSON file.
        segmented_json_file (str): Path to consultation paper segmented JSON file.
        output_dir (str): Directory where the generated reports will be saved.
    """
    main_word(
        consult_quest_summary_json,
        consult_paper_info_json,
        executive_summary_json,
        output_dir,
    )
