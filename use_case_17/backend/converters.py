from io import BytesIO
from pathlib import Path
import tempfile
import os


def _md_to_html(md: str) -> str:
    try:
        import markdown as _md

        return _md.markdown(md, extensions=["extra"])
    except ImportError:
        import pypandoc

        return pypandoc.convert_text(md, to="html", format="md")


def _pandoc_md_to_docx(md: str) -> bytes:
    import pypandoc

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_name = tmp.name
    try:
        pypandoc.convert_text(md, to="docx", format="md", outputfile=tmp_name)
        return Path(tmp_name).read_bytes()
    finally:
        try:
            os.remove(tmp_name)
        except OSError:
            pass


def _fallback_md_to_docx(md: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    style_cfg = {
        "Normal": 11,
        "Heading 1": 16,
        "Heading 2": 14,
    }
    for style_name, size in style_cfg.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(3)

    for line in md.splitlines():
        line = line.rstrip()
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        else:
            doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

def _force_black_font(doc_bytes: bytes) -> bytes:
    from docx import Document
    from docx.shared import RGBColor

    doc = Document(BytesIO(doc_bytes))
    black = RGBColor(0, 0, 0)

    for p in doc.paragraphs:
        for run in p.runs:
            run.font.color.rgb = black

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = black

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

def md_to_docx(md: str) -> bytes:
    try:
        doc_bytes = _pandoc_md_to_docx(md)
    except Exception:
        doc_bytes = _fallback_md_to_docx(md)
    return _force_black_font(doc_bytes)


def md_to_pdf(md: str) -> bytes:
    from weasyprint import HTML, CSS

    html = _md_to_html(md)

    calibri_css = CSS(
        string="""
        @page { margin: 2cm; }
        body, h1, h2, h3, h4, h5, h6, p {
            font-family: 'Calibri', sans-serif;
            color: #000000;
            margin: 0 0 6pt 0;   /* single consistent spacing */
        }
        """
    )

    return HTML(string=html, base_url=str(Path(".").resolve())).write_pdf(
        stylesheets=[calibri_css]
    )
