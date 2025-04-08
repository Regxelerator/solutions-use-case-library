import os
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from utils.file_handler import read_json_file


def set_run_font(
    run, font_name="Calibri", font_size=10, is_bold=False, font_color="000000"
):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = is_bold
    run.font.color.rgb = RGBColor.from_string(font_color)


def set_heading_run_font(
    run, font_name="Calibri Light", font_size=14, is_bold=True, font_color="000000"
):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = is_bold
    run.font.color.rgb = RGBColor.from_string(font_color)


def create_footer_disclaimer_and_page_number(section):
    footer = section.footer
    footer.paragraphs.clear()
    para_left = footer.add_paragraph()
    para_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_left = para_left.add_run(
        "Disclaimer: This output was created through the application of generative AI. "
        "Please validate accuracy and completeness before using it for decision-making."
    )
    set_run_font(run_left, font_size=8, font_name="Calibri")
    para_right = footer.add_paragraph()
    para_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_right = para_right.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run_right._r.append(fldChar_begin)
    run_right._r.append(instrText)
    run_right._r.append(fldChar_separate)
    run_right._r.append(fldChar_end)
    set_run_font(run_right, font_size=8, font_name="Calibri")


def set_cell_shading(cell, fill_color):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_borders(cell, border_color="000000", border_size="4"):
    tcPr = cell._element.get_or_add_tcPr()
    for border_tag in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(border_tag)
        border_elm = OxmlElement(tag)
        border_elm.set(qn("w:val"), "single")
        border_elm.set(qn("w:sz"), border_size)
        border_elm.set(qn("w:color"), border_color)
        border_elm.set(qn("w:space"), "0")
        tcPr.append(border_elm)


def extract_authority_from_file(path):
    """
    Extracts the 'authority' field from the 'regulation_overview' section of a JSON file.

    Parameters:
    - path (str): The file path to the JSON file containing the benchmarking or regulatory data.

    Returns:
    - str: The 'authority' value from the 'regulation_overview' section if found, otherwise an empty string.

    """
    if not os.path.exists(path):
        return ""

    raw_data = read_json_file(path)
    if "regulation_overview" not in raw_data:
        return ""
    overview_str = raw_data["regulation_overview"]
    try:
        parsed_overview = json.loads(overview_str)
    except Exception:
        return ""
    items = parsed_overview.get("regulation_overview", [])
    if not items:
        return ""
    return items[0].get("authority", "").strip()


def generating_benchmarking_report(input_dir, final_analysis_output_path):
    """
    Generates a benchmarking report based on generated benchmarking report

    Parameters:
    - input_dir (str): The directory containing the input data files (e.g., JSON files) for benchmarking analysis.
    - final_analysis_output_path (str): The file path where the generated benchmarking report will be saved.

    Returns:
    str : path the generated report to the specified output path.

    """
    if not final_analysis_output_path:
        print("Error: Comparative analysis failed - cannot generate report")
        return None

    data = read_json_file(final_analysis_output_path)

    processed_dir = os.path.join(input_dir, "Regulations", "Processed Regulations")
    files = ["regulation_1.json", "regulation_2.json"]
    authorities = []
    for filename in files:
        full_path = os.path.join(processed_dir, filename)
        authority = extract_authority_from_file(full_path)
        authorities.append(authority)

    def valid_provisions(provisions_dict):
        vals = []
        for v in provisions_dict.values():
            if v.strip() and "No data provided" not in v:
                vals.append(v)
        return vals

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    create_footer_disclaimer_and_page_number(section)
    title_par = doc.add_paragraph()
    title_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = title_par.add_run("Benchmarking analysis results")
    set_heading_run_font(
        run_title, font_size=20, font_name="Calibri Light", is_bold=True
    )
    doc.add_paragraph()
    max_countries = 0
    for item in data["benchmarking_analysis"]:
        c_provisions = valid_provisions(item.get("country_provisions", {}))
        if len(c_provisions) > max_countries:
            max_countries = len(c_provisions)
    table = doc.add_table(rows=1, cols=max_countries + 1)
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for i in range(max_countries):
        if i < len(authorities) and authorities[i]:
            hdr_text = f"Provisions ({authorities[i]})"
        else:
            hdr_text = "Provisions"
        hdr_cells[i].text = ""
        r = hdr_cells[i].paragraphs[0].add_run(hdr_text)
        set_run_font(
            r, font_name="Calibri", font_size=10, is_bold=True, font_color="FFFFFF"
        )
        set_cell_shading(hdr_cells[i], "000000")
        set_cell_borders(hdr_cells[i])
    hdr_cells[max_countries].text = ""
    r_ca = hdr_cells[max_countries].paragraphs[0].add_run("Comparative Analysis")
    set_run_font(
        r_ca, font_name="Calibri", font_size=10, is_bold=True, font_color="FFFFFF"
    )
    set_cell_shading(hdr_cells[max_countries], "000000")
    set_cell_borders(hdr_cells[max_countries])
    for item in data["benchmarking_analysis"]:
        dim_row = table.add_row()
        dim_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        dim_row.height = Inches(0.4)
        for c in dim_row.cells:
            c.text = ""
            set_cell_borders(c)
        dim_cell = dim_row.cells[0]
        dim_cell.merge(dim_row.cells[max_countries])
        run_dim = dim_cell.paragraphs[0].add_run(item["dimension_name"])
        set_run_font(run_dim, font_name="Calibri", font_size=10, is_bold=True)
        set_cell_shading(dim_cell, "D9D9D9")
        val_row = table.add_row()
        for c in val_row.cells:
            set_cell_borders(c)
        c_provisions = valid_provisions(item.get("country_provisions", {}))
        for idx, provision in enumerate(c_provisions):
            run_cp = val_row.cells[idx].paragraphs[0].add_run(provision)
            set_run_font(run_cp, font_name="Calibri", font_size=10)
        run_comp = (
            val_row.cells[max_countries]
            .paragraphs[0]
            .add_run(item.get("comparative_analysis", ""))
        )
        set_run_font(run_comp, font_name="Calibri", font_size=10)

    output_dir = os.path.join(os.getcwd(), "output")
    os.makedirs(output_dir, exist_ok=True)
    Final_report_path = os.path.join(output_dir, "Benchmarking_Analysis_Report.docx")
    doc.save(Final_report_path)
    return Final_report_path
