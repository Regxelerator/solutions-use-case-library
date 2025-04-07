import docx
import PyPDF2
import os


def extract_text_from_docx(file_path: str) -> str:
    """Extracts text from a DOCX file.

    :param file_path: Path to the DOCX file.
    :return: Extracted text as a string or an error message.
    """
    try:
        print(f"data is being read from doxs.........")
        document = docx.Document(file_path)
        print("\n".join([para.text for para in document.paragraphs]))
        return "\n".join([para.text for para in document.paragraphs])
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}"


def extract_text_from_pdf(file_path: str) -> str:
    """Extracts text from a PDF file.

    :param file_path: Path to the PDF file.
    :return: Extracted text as a string or an error message.
    """
    text = ""
    try:
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            print(f"data is being read from pdf.........")
            for page in reader.pages:
                page_text = page.extract_text()
                print(f"\n{page_text}.")
                if page_text:
                    text += page_text + "\n"

        return text.strip()
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}"


def identify_file_type(file_path: str) -> str:
    """Identifies the file type based on its extension.

    :param file_path: Path to the file.
    :return: File extension in lowercase without the leading dot.
    """
    _, extension = os.path.splitext(file_path)
    return extension.lower().strip(".")
