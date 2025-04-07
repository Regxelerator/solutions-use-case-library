from datetime import datetime
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import _Cell, Table
import pandas as pd
import json
import os


def set_run_font(
    run: Run, font_name: str = "Calibri", font_size: int = 10, is_bold: bool = False
) -> None:
    """
    Sets the font properties of a Word document run.

    Args:
        run (Run): The text run whose font properties will be modified.
        font_name (str, optional): The name of the font. Defaults to "Calibri".
        font_size (int, optional): The font size in points. Defaults to 10.
        is_bold (bool, optional): Whether the text should be bold. Defaults to False.
    """
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = is_bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_cell_font(
    cell: _Cell,
    text: str,
    font_name: str = "Calibri",
    font_size: int = 10,
    bold: bool = False,
) -> None:
    """
    Sets the font properties for text inside a Word table cell.

    Args:
        cell (_Cell): The table cell where text formatting will be applied.
        text (str): The text to be added to the cell.
        font_name (str, optional): The name of the font. Defaults to "Calibri".
        font_size (int, optional): The font size in points. Defaults to 10.
        bold (bool, optional): Whether the text should be bold. Defaults to False.
    """
    cell.text = ""
    run: Run = cell.paragraphs[0].add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold


def set_cell_borders(
    cell: _Cell, border_color: str = "000000", border_size: int = 4
) -> None:
    """
    Sets the border properties for a table cell in a Word document.

    Args:
        cell (_Cell): The table cell where the border will be applied.
        border_color (str, optional): The hex color code of the border. Defaults to "000000" (black).
        border_size (int, optional): The thickness of the border in eighths of a point. Defaults to 4.
    """
    tcPr = cell._element.get_or_add_tcPr()
    for border_tag in ("top", "left", "bottom", "right"):
        tag = f"w:{border_tag}"
        border_elm = OxmlElement(tag)
        border_elm.set(qn("w:val"), "single")
        border_elm.set(qn("w:sz"), str(border_size))
        border_elm.set(qn("w:color"), border_color)
        border_elm.set(qn("w:space"), "0")
        tcPr.append(border_elm)


def _insert_bullet_points(doc: Document, gpt_markdown_text: str) -> None:
    """
    Inserts bullet points into a Word document from a given markdown-style text.

    Args:
        doc (Document): The Word document where the bullet points will be added.
        gpt_markdown_text (str): Markdown-style text containing bullet points.
    """
    if isinstance(gpt_markdown_text, dict):
        gpt_markdown_text = json.dumps(gpt_markdown_text)
    for line in gpt_markdown_text.splitlines():
        line = line.strip()
        if not line:
            continue

        if line.startswith("-"):
            paragraph: Paragraph = doc.add_paragraph(style="List Bullet")
            line = line.lstrip("-").strip()
        else:
            paragraph: Paragraph = doc.add_paragraph()

        parts = re.split(r"(\*\*.*?\*\*)", line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                bold_text: str = part[2:-2]
                run: Run = paragraph.add_run(bold_text)
                set_run_font(run, is_bold=True)
            else:
                run: Run = paragraph.add_run(part)
                set_run_font(run, is_bold=False)


def _create_pandas_table_in_doc(doc: Document, df: pd.DataFrame) -> None:
    """
    Creates and inserts a table into a Word document.

    Args:
        doc (Document): The Word document where the table will be added.
        df (pd.DataFrame): The DataFrame containing table data.
    """
    df_rows = [df.columns.tolist()] + df.values.tolist()
    table: Table = doc.add_table(rows=len(df_rows), cols=len(df_rows[0]))
    table.style = "Table Grid"

    for i, row in enumerate(df_rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            set_cell_font(cell, str(val), bold=(i == 0))
            set_cell_borders(cell)

    doc.add_paragraph()


def create_board_meeting_analysis_report(
    df_director_attendance: pd.DataFrame,
    df_position_matrix: pd.DataFrame,
    analysis_attendance_text: str,
    df_items_per_meeting: pd.DataFrame,
    df_by_category: pd.DataFrame,
    analysis_effectiveness_text: str,
    output_folder_path: str,
) -> None:
    """
    Generates a Word document report analyzing board meeting attendance, agenda items, and effectiveness.

    Args:
        df_director_attendance (pd.DataFrame): DataFrame containing attendance data for board directors.
        df_position_matrix (pd.DataFrame): DataFrame representing the position breakdown per meeting.
        analysis_attendance_text (str): Text containing analysis of attendance in bullet points format.
        df_items_per_meeting (pd.DataFrame): DataFrame summarizing the number of agenda items per meeting.
        df_by_category (pd.DataFrame): DataFrame categorizing agenda items across meetings.
        analysis_effectiveness_text (str): Text containing analysis of meeting effectiveness in bullet points format.
        output_folder_path (str) : path to file to save report
    Returns:
        None: Saves the generated report as a Word document to given path.
    """

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    doc.styles["Heading 1"].font.name = "Calibri Light"
    doc.styles["Heading 2"].font.name = "Calibri Light"

    doc.styles["List Bullet"].font.name = "Calibri"
    doc.styles["List Bullet"].font.size = Pt(10)

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(1)

    heading_paragraph = doc.add_heading("Results – Board meeting analysis", level=1)
    set_run_font(
        heading_paragraph.runs[0], font_size=16, font_name="Calibri Light", is_bold=True
    )
    heading_paragraph.paragraph_format.space_after = Pt(12)

    today_str = datetime.now().strftime("%Y-%m-%d")
    details = {
        "Date of assessment": today_str,
        "Entity": "Placeholder Entity Name",
        "Time period covered": "January 2024 - December 2024",
    }

    for key, value in details.items():
        paragraph = doc.add_paragraph()
        key_run = paragraph.add_run(f"{key}: ")
        set_run_font(key_run, is_bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run)
        paragraph.paragraph_format.space_after = Pt(4)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    section_heading = doc.add_heading(
        "1. Analysis - Meeting attendee analysis", level=2
    )
    set_run_font(
        section_heading.runs[0], font_size=12, font_name="Calibri Light", is_bold=False
    )
    section_heading.paragraph_format.space_after = Pt(8)

    doc.add_paragraph("Attendance summary of Board Directors:")
    _create_pandas_table_in_doc(doc, df_director_attendance)

    doc.add_paragraph("Position category breakdown by meeting:")
    _create_pandas_table_in_doc(doc, df_position_matrix)

    doc.add_paragraph("Key observations:").paragraph_format.space_after = Pt(4)
    _insert_bullet_points(doc, analysis_attendance_text)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    section_heading = doc.add_heading("2. Analysis - Meeting agenda items", level=2)
    set_run_font(
        section_heading.runs[0], font_size=12, font_name="Calibri Light", is_bold=False
    )
    section_heading.paragraph_format.space_after = Pt(8)

    doc.add_paragraph("Number of agenda items per meeting:")
    _create_pandas_table_in_doc(doc, df_items_per_meeting)

    doc.add_paragraph(
        "Breakdown of Agenda Items by Category (including All Meetings Combined):"
    )
    _create_pandas_table_in_doc(doc, df_by_category)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    section_heading = doc.add_heading(
        "3. Analysis - Ancillary meeting effectiveness criteria", level=2
    )
    set_run_font(
        section_heading.runs[0], font_size=12, font_name="Calibri Light", is_bold=False
    )
    section_heading.paragraph_format.space_after = Pt(8)

    _insert_bullet_points(doc, analysis_effectiveness_text)

    footer = section.footer
    footer_paragraph = (
        footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    )
    footer_paragraph.clear()
    run = footer_paragraph.add_run(
        "Disclaimer: This output was created through the application of generative AI. "
        "Please validate accuracy and completeness before using it for decision-making."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    file_path = os.path.join(output_folder_path, "Board_Meeting_Analysis_Report.docx")
    doc.save(file_path)
    print(f"Board meeting analysis report created successfully: {file_path}")
