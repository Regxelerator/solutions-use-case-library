from __future__ import annotations

import argparse
import json
from datetime import datetime
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _load_json(path: Path | str):
    with open(path, encoding="utf-8") as fp:
        return json.load(fp)


def _collect_registrant_names(filings_dir: Path) -> List[str]:
    names: set[str] = set()
    for p in filings_dir.rglob("*.json"):
        with p.open("r", encoding="utf-8") as fp:
            data = json.load(fp)
        names.add((data.get("EntityRegistrantName") or p.stem).strip())
    return sorted(names)


def _set_run_font(run, *, size=11, bold=False, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def _set_para_spacing(p, *, after=0, before=0):
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)


def _add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    _set_run_font(h.runs[0], size=14 if level == 1 else 12, bold=(level == 1), name="Calibri Light")
    if level == 1:
        h.paragraph_format.space_after = Pt(12)
    else:
        _set_para_spacing(h)

def _add_footer(section):
    p = section.footer.add_paragraph()
    _set_para_spacing(p)
    _set_run_font(p.add_run("Disclaimer: This output was created through the application of generative AI. Please validate accuracy and completeness before using it for decision-making."), size=8)


def _add_shaded_para(doc: Document, text: str):
    p = doc.add_paragraph()
    _set_para_spacing(p)
    _set_run_font(p.add_run(text), size=10)
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shd)


def _add_bullet_or_para(doc: Document, line: str):
    cleaned = line.lstrip("•- ").strip()
    if not cleaned:
        return
    style = "List Bullet" if line.lstrip().startswith(("•", "-")) else None
    p = doc.add_paragraph(style=style)
    _set_para_spacing(p)
    _set_run_font(p.add_run(cleaned), size=10)


def _render_block(doc: Document, text: str):
    for ln in text.splitlines():
        _add_bullet_or_para(doc, ln)


def build_report(analysis: Dict, registrants: List[str], output_path: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(1)
    sec.top_margin = sec.bottom_margin = Inches(1)
    _add_footer(sec)

    tp = doc.add_paragraph(); _set_para_spacing(tp)
    _set_run_font(tp.add_run("Cybersecurity Disclosure Comparative Assessment"), size=20, bold=True, name="Calibri Light")
    dp = doc.add_paragraph(); _set_para_spacing(dp)
    _set_run_font(dp.add_run(datetime.now().strftime("%d %B %Y")), size=11)
    doc.add_paragraph()

    parents = sorted(r for r,v in analysis.items() if v.get("hierarchy_level") == 1)
    for parent_id in parents:
        parent = analysis[parent_id]
        _add_heading(doc, f"{parent_id} – {parent['requirement']}", level=1)
        es = parent.get("executive_summary", {}).get("executive_summary", "")
        if es:
            _add_shaded_para(doc, es)
            doc.add_paragraph()

        related = sorted(r for r in analysis if r==parent_id or r.startswith(f"{parent_id}."))
        for rid in related:
            entry = analysis[rid]
            lvl = entry.get("hierarchy_level")
            if lvl in (2,3):
                _add_heading(doc, rid, level=2)
                req_p = doc.add_paragraph(); _set_para_spacing(req_p)
                run_req = req_p.add_run(entry["requirement"])
                run_req.italic = True
                _set_run_font(run_req, size=10)
                req_p.paragraph_format.space_after = Pt(8)
            else:
                _add_heading(doc, f"{rid} – {entry['requirement']}", level=2)

            ar = (entry.get("analysis") or {}).get("analysis_results", {})
            if not ar:
                p = doc.add_paragraph("No disclosures available."); _set_para_spacing(p); _set_run_font(p.runs[0], size=10)
                doc.add_paragraph(); continue

            _add_heading(doc, "Key trends and commonalities", level=3)
            _render_block(doc, ar.get("key_trends_and_commonalities", ""))
            spacer = doc.add_paragraph(); spacer.paragraph_format.space_after = Pt(4)
            _add_heading(doc, "Key differences and notable outliers", level=3)
            _render_block(doc, ar.get("key_differences_and_notable_outliers", ""))
            doc.add_paragraph()
        doc.add_page_break()

    _add_heading(doc, "Appendix 1 – Entities in scope of the analysis", level=1)
    for n in registrants:
        p = doc.add_paragraph(n); _set_para_spacing(p); _set_run_font(p.runs[0], size=10)

    doc.save(output_path); print(f"Report saved: {output_path}")


def main():
    ap = argparse.ArgumentParser(description="Generate Word report from analysis results")
    ap.add_argument("--analysis-path", type=Path, default=Path("output/analysis_results.json"))
    ap.add_argument("--filings-dir",   type=Path, default=Path("output/entity_filings_json"))
    ap.add_argument("--output-path",   type=Path, default=Path("output/Cybersecurity_Disclosure_Assessment_Results.docx"))
    args = ap.parse_args()

    build_report(_load_json(args.analysis_path), _collect_registrant_names(args.filings_dir), args.output_path)

if __name__ == "__main__":
    main()