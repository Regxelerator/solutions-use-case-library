import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
import sys
from utils.file_handler import load_json_file


def set_run_font(run, font_name="Calibri", font_size=10, is_bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = is_bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_cell_font(cell, text, font_name="Calibri", font_size=10, bold=False):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold


def set_cell_borders(cell, border_color="000000", border_size=4):
    tcPr = cell._element.get_or_add_tcPr()
    for border_tag in ("top", "left", "bottom", "right"):
        tag = f"w:{border_tag}"
        border_elm = OxmlElement(tag)
        border_elm.set(qn("w:val"), "single")
        border_elm.set(qn("w:sz"), str(border_size))
        border_elm.set(qn("w:color"), border_color)
        border_elm.set(qn("w:space"), "0")
        tcPr.append(border_elm)


def set_cell_shading(cell, fill_color):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def get_color_for_rating(rating_str):
    if str(rating_str).strip().upper() == "NOT APPLICABLE":
        return "#D9D9D9"
    try:
        val = float(rating_str)
    except ValueError:
        return "#D9D9D9"
    if val <= 3:
        return "#C00000"
    elif val <= 6:
        return "#FFD966"
    elif val <= 8:
        return "#D6E3BC"
    else:
        return "#A8D08D"


def create_assessment_table(doc, list_of_items):
    table = doc.add_table(rows=1, cols=4)

    header_cells = table.rows[0].cells
    header_cells[0].text = "ID"
    header_cells[1].text = "Requirement"
    header_cells[2].text = "Compliance Assessment"
    header_cells[3].text = "Compliance score"

    for cell in header_cells:
        set_run_font(
            cell.paragraphs[0].runs[0], font_name="Calibri", font_size=11, is_bold=True
        )
        set_cell_shading(cell, "0E2841")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_borders(cell, border_color="000000")

    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)

    for item in list_of_items:
        row = table.add_row().cells

        row[0].text = item.get("id", "")
        row[1].text = item.get("requirement_description", "")
        row[2].text = item.get("compliance_assessment", "")
        row[3].text = str(item.get("compliance_score", ""))

        for c in row:
            set_cell_borders(c, border_color="000000")
            if c.paragraphs[0].runs:
                set_run_font(c.paragraphs[0].runs[0], font_name="Calibri", font_size=10)

        rating_str = row[3].text
        fill_color = get_color_for_rating(rating_str).lstrip("#")
        set_cell_shading(row[3], fill_color)

    for row in table.rows:
        row.cells[0].width = Cm(1.0)
        row.cells[1].width = Cm(6.8)
        row.cells[2].width = Cm(6.8)
        row.cells[3].width = Cm(1.8)

    return table


def create_word_document_from_json(list_data, output_file_location):
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    heading_paragraph = doc.add_heading(
        "Results – Outsourcing Contract Assessment", level=1
    )
    set_run_font(
        heading_paragraph.runs[0], is_bold=True, font_size=18, font_name="Calibri Light"
    )
    heading_paragraph.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()

    create_assessment_table(doc, list_data)

    section = doc.sections[0]
    footer = section.footer
    footer_paragraph = (
        footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    )
    footer_paragraph.clear()
    run = footer_paragraph.add_run(
        "Disclaimer: This output was created through the application of generative AI. "
        "Please verify its accuracy before using it for decisions."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(8)

    doc.save(output_file_location)
    print(f"Document created successfully at {output_file_location}")
    return output_file_location


def create_final_report(contract_assessment_json_path: str, output_dir: str) -> str:
    """
    Creates a Word document report based on the outsourcing contract compliance assessment data.

    Args:
        contract_assessment_json_path (str):
            Path to the JSON file containing the results of the outsourcing contract compliance assessment.
        output_dir (str):
            Path to the directory where the final Word report will be saved.

    Returns:
        str:
            The full path of the generated Word document."
    """
    json_data = load_json_file(contract_assessment_json_path)
    if json_data is None:
        sys.exit("Error: Could not load the JSON data (did not find or parse).")
    final_report_path = os.path.join(output_dir, "Final Report")
    os.makedirs(final_report_path, exist_ok=True)
    output_docx_filename = "Results_Outsourcing_Contract_Assessment.docx"
    output_file_location = os.path.join(final_report_path, output_docx_filename)
    output_path = create_word_document_from_json(json_data, output_file_location)
    print(f"Document created at: {output_path}")
    return output_path
