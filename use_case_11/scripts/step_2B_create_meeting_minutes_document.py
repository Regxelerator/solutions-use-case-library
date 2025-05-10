from __future__ import annotations
import json
import os
import re
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

def set_run_font(
    run,
    *,
    font: str = "Calibri",
    size: int = 10,
    bold: bool = False,
    color: tuple[int, int, int] = (0, 0, 0),
) -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)

def set_heading_run_font(run, *, font: str = "Calibri Light", size: int = 14) -> None:
    set_run_font(run, font=font, size=size, bold=True)

def add_footer(section, text: str) -> None:
    footer = section.footer
    para = footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    set_run_font(run, size=8)

def build_word(doc_data: dict, out_path: os.PathLike) -> None:
    meeting_dt = datetime.strptime(doc_data["meeting_date"], "%Y-%m-%d")
    date_str = f"{meeting_dt.day} {meeting_dt.strftime('%B %Y')}"

    doc = Document()
    cover = doc.sections[0]
    cover.top_margin = Inches(4)
    cover.bottom_margin = Inches(3)
    cover.left_margin = cover.right_margin = Inches(1)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = (
        "Monetary Policy Summary and minutes of the Monetary Policy Committee\n"
        f"meeting ending on {date_str}"
    )
    title_run = title_para.add_run(title)
    set_heading_run_font(title_run, size=28)

    doc.add_page_break()

    body_sec = doc.add_section()
    body_sec.top_margin = Inches(1)
    body_sec.bottom_margin = Inches(1.5)
    body_sec.left_margin = body_sec.right_margin = Inches(1)
    add_footer(
        body_sec,
        f"Bank of England Minutes of the Monetary Policy Committee "
        f"Meeting ending on {date_str}",
    )

    para_counter = 1
    for section in doc_data["sections"]:
        name = section["section_name"]
        paragraphs = section.get("section_paragraphs", [])

        if name not in {"Introduction", "Attendees"}:
            hdr = doc.add_paragraph()
            hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_hdr = hdr.add_run(name)
            set_run_font(run_hdr, bold=True)

        for raw in paragraphs:
            cleaned = re.sub(
                r"\*\*Bank of England Minutes of the Monetary Policy Committee "
                r"Meeting.*?\*\*\s*\d*",
                "",
                raw,
                flags=re.I,
            ).rstrip()

            if not cleaned:
                continue

            if name != "Attendees":
                cleaned = re.sub(r"\s*\n\s*", " ", cleaned)
                para = doc.add_paragraph()
                run = para.add_run(f"{para_counter} {cleaned}")
                set_run_font(run)
                para.paragraph_format.space_after = Pt(6)

            else:
                lines = [ln.strip() for ln in cleaned.split("\n") if ln.strip()]
                para = doc.add_paragraph()
                first = para.add_run(f"{para_counter} {lines[0]}")
                set_run_font(first)
                for ln in lines[1:]:
                    para.add_run().add_break()
                    run_ln = para.add_run(ln)
                    set_run_font(run_ln)
                para.paragraph_format.space_after = Pt(6)

            para_counter += 1

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"✓ Word document created: {out_path}")


def generate_word_minutes(dataset_dir: str | os.PathLike = "output/test_dataset") -> None:
    dataset_dir = Path(dataset_dir)
    if not dataset_dir.is_dir():
        raise FileNotFoundError(dataset_dir)

    for json_path in dataset_dir.glob("*.json"):
        with open(json_path, encoding="utf-8") as f:
            data = json.load(f)

        meeting_dt = datetime.strptime(data["meeting_date"], "%Y-%m-%d")
        stamp = meeting_dt.strftime("%Y_%m_%d")
        doc_name = f"MPC_Minutes_{stamp}.docx"
        build_word(data, dataset_dir / doc_name)